"""
Word文档生成模块
用于基于模板生成Word报告
"""
from docx import Document
from datetime import datetime
import re


class WordGenerator:
    """Word文档生成器"""
    
    def __init__(self, template_path):
        """
        初始化Word生成器
        
        Args:
            template_path: 模板文件路径
        """
        self.template_path = template_path
    
    def generate(self, data, output_path):
        """
        生成Word文档
        
        Args:
            data: 包含统计数据的字典
            output_path: 输出文件路径
        
        Returns:
            bool: 是否成功生成
        """
        try:
            # 读取模板
            doc = Document(self.template_path)
            
            # 准备替换数据
            replacements = {
                # 基础统计
                '{{total_count}}': str(data.get('total_current', 0)),
                '{{sunshine_count}}': str(data.get('sunshine_current', 0)),
                '{{last_week_sunshine}}': str(data.get('sunshine_last', 0)),
                '{{sunshine_trend}}': data.get('sunshine_trend', '持平'),
                '{{gab_count}}': str(data.get('gab_current', 0)),
                '{{last_week_gab}}': str(data.get('gab_last', 0)),
                '{{gab_trend}}': data.get('gab_trend', '持平'),
                
                # 人员信息
                '{{sunshine_persons}}': data.get('sunshine_persons_text', ''),
                '{{gab_persons}}': data.get('gab_persons_text', ''),
                
                # 统计分析
                '{{area_stats}}': data.get('area_stats_text', ''),
                '{{group_appeal}}': data.get('group_appeal_text', '无'),
                '{{travel_road_count}}': str(data.get('travel_road_count', 0)),
                '{{travel_stats}}': data.get('travel_stats_text', '无'),
            }
            
            # 遍历所有段落，替换占位符
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        # 替换段落中的占位符
                        self._replace_text_in_paragraph(paragraph, placeholder, value)
            
            # 遍历所有表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacements.items():
                                if placeholder in paragraph.text:
                                    self._replace_text_in_paragraph(paragraph, placeholder, value)
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"生成Word文档失败: {e}")
            return False
    
    def _replace_text_in_paragraph(self, paragraph, placeholder, value):
        """
        在段落中替换占位符，完整保留格式（字体、缩进等）
        
        Args:
            paragraph: 段落对象
            placeholder: 占位符
            value: 替换值
        """
        # 检查占位符是否在段落文本中
        if placeholder not in paragraph.text:
            return
        
        # 遍历段落中的所有run，在每个run中替换文本
        for run in paragraph.runs:
            if placeholder in run.text:
                # 直接替换run中的文本，保留所有格式
                run.text = run.text.replace(placeholder, value)


def test_generator():
    """测试Word生成器"""
    print("=== Word生成模块测试 ===\n")
    
    # 准备测试数据
    test_data = {
        'total_current': 5,
        'sunshine_current': 2,
        'sunshine_last': 2,
        'sunshine_trend': '持平',
        'gab_current': 3,
        'gab_last': 2,
        'gab_trend': '上升1人',
    }
    
    # 生成文档
    generator = WordGenerator('template.docx')
    output_file = 'output/测试报告.docx'
    
    success = generator.generate(test_data, output_file)
    
    if success:
        print(f"✓ Word文档生成成功: {output_file}")
        print(f"\n数据内容:")
        print(f"  本周总计: {test_data['total_current']}人")
        print(f"  阳光xf登记: {test_data['sunshine_current']}人 (上周{test_data['sunshine_last']}人，{test_data['sunshine_trend']})")
        print(f"  gab上访: {test_data['gab_current']}人 (上周{test_data['gab_last']}人，{test_data['gab_trend']})")
    else:
        print("❌ Word文档生成失败")


if __name__ == '__main__':
    test_generator()


